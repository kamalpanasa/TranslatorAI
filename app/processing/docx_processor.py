import os
from typing import Callable, Coroutine, Any
from docx import Document
from app.utils.logger import get_logger
from app.utils.exceptions import FileProcessingError

logger = get_logger("app.processing.docx_processor")


class DocxProcessor:
    @staticmethod
    async def translate_docx_file(
        input_path: str,
        output_path: str,
        translate_func: Callable[[str], Coroutine[Any, Any, str]]
    ) -> str:
        """
        Translates a DOCX file:
        1. Open the original Word document.
        2. Translate body paragraphs.
        3. Translate table cell paragraphs.
        4. Save document.
        """
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"DOCX not found: {input_path}")

        try:
            logger.info(f"Opening DOCX for translation: {input_path}")
            doc = Document(input_path)
            
            # Translate body paragraphs
            for para in doc.paragraphs:
                clean_text = para.text.strip()
                if clean_text:
                    # Translate paragraph text
                    translated_text = await translate_func(clean_text)
                    # Set text while keeping paragraph format
                    para.text = translated_text

            # Translate table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Translate paragraphs inside table cell
                        for para in cell.paragraphs:
                            clean_text = para.text.strip()
                            if clean_text:
                                translated_text = await translate_func(clean_text)
                                para.text = translated_text

            # Save the modified document
            doc.save(output_path)
            logger.info(f"Translated DOCX saved successfully to: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error processing DOCX document: {str(e)}", exc_info=True)
            raise FileProcessingError(f"Failed to translate DOCX document: {str(e)}")
